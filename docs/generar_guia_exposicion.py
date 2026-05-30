"""
Genera la guia de exposicion en Word (.docx) para el Grupo 6.
Salida: docs/Guia-Exposicion-Backend-Intranet.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


COLOR_PRIMARIO = RGBColor(0x21, 0x4F, 0x99)
COLOR_GRIS = RGBColor(0x55, 0x55, 0x55)
COLOR_VERDE = RGBColor(0x2E, 0x7D, 0x32)
COLOR_NEGRO = RGBColor(0x00, 0x00, 0x00)


def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def estilo_titulo(p, texto, size=18, color=COLOR_PRIMARIO, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT):
    p.alignment = align
    run = p.add_run(texto)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def parrafo(doc, texto, size=11, bold=False, italic=False, color=COLOR_NEGRO, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(texto)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return p


def heading(doc, texto, nivel=1):
    sizes = {1: 18, 2: 14, 3: 12}
    colors = {1: COLOR_PRIMARIO, 2: COLOR_PRIMARIO, 3: COLOR_GRIS}
    p = doc.add_paragraph()
    estilo_titulo(p, texto, size=sizes.get(nivel, 11), color=colors.get(nivel, COLOR_NEGRO))


def codigo(doc, texto, lenguaje="java"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(texto)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x10, 0x10, 0x10)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F4F4")
    p_pr.append(shd)


def bullet(doc, texto, nivel=0):
    p = doc.add_paragraph(style="List Bullet")
    if nivel > 0:
        p.paragraph_format.left_indent = Cm(0.5 * (nivel + 1))
    run = p.runs[0] if p.runs else p.add_run(texto)
    if not p.runs:
        run.text = texto
    else:
        run.text = texto
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def tabla_simple(doc, headers, rows, anchos=None):
    tabla = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tabla.style = "Light Grid Accent 1"

    hdr = tabla.rows[0]
    for i, texto in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(texto)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(cell, "214F99")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i, fila in enumerate(rows):
        row = tabla.rows[i + 1]
        for j, texto in enumerate(fila):
            cell = row.cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(texto))
            run.font.size = Pt(9)

    if anchos:
        for row in tabla.rows:
            for i, cell in enumerate(row.cells):
                if i < len(anchos):
                    cell.width = Cm(anchos[i])

    return tabla


def callout(doc, titulo, texto, color_hex="FFF4E1"):
    tabla = doc.add_table(rows=1, cols=1)
    cell = tabla.rows[0].cells[0]
    set_cell_background(cell, color_hex)
    cell.text = ""
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(f"💡 {titulo}\n")
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = COLOR_PRIMARIO
    p2 = cell.add_paragraph()
    r2 = p2.add_run(texto)
    r2.font.size = Pt(10)


def saltopagina(doc):
    doc.add_page_break()


def build_doc():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ═══════════════════════════════════════════════════════════════
    # PORTADA
    # ═══════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    estilo_titulo(p, "GUÍA DE EXPOSICIÓN", size=26, color=COLOR_PRIMARIO,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    estilo_titulo(p, "Backend Intranet TelecoPerú", size=20, color=COLOR_GRIS,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    estilo_titulo(p, "Sistema de Gestión de Incidencias", size=14, color=COLOR_GRIS,
                  bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    doc.add_paragraph()

    parrafo(doc, "Curso: Desarrollo Web Integrado", size=12, bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER)
    parrafo(doc, "Docente: Marcelino Estrada Aro", size=11,
            align=WD_ALIGN_PARAGRAPH.CENTER)
    parrafo(doc, "Grupo 6 — Unidad 2", size=11,
            align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    doc.add_paragraph()

    parrafo(doc, "INTEGRANTES", size=13, bold=True, color=COLOR_PRIMARIO,
            align=WD_ALIGN_PARAGRAPH.CENTER)

    integrantes_tabla = doc.add_table(rows=5, cols=2)
    integrantes_tabla.alignment = WD_ALIGN_PARAGRAPH.CENTER
    integrantes = [
        ("GONZALES ALVIS, Claudia Leonor", "Development Team — Frontend Angular"),
        ("PRADO MISAICO, Bartolome Angelo", "Development Team — Backend"),
        ("RODRIGUEZ POZO, Matias Ariel", "Product Owner"),
        ("HUANE SARMIENTO, Carlos Jesús", "Scrum Master"),
        ("RODRIGUEZ CHACALIAZA, Airton Clides", "Development Team — BD / QA"),
    ]
    for i, (nom, rol) in enumerate(integrantes):
        c1 = integrantes_tabla.rows[i].cells[0]
        c2 = integrantes_tabla.rows[i].cells[1]
        c1.text = ""
        c2.text = ""
        r1 = c1.paragraphs[0].add_run(nom)
        r1.font.size = Pt(11)
        r1.font.bold = True
        r2 = c2.paragraphs[0].add_run(rol)
        r2.font.size = Pt(11)
        r2.font.italic = True
        r2.font.color.rgb = COLOR_GRIS

    saltopagina(doc)

    # ═══════════════════════════════════════════════════════════════
    # ÍNDICE
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "Índice", nivel=1)
    indice = [
        ("1. Cómo usar esta guía", "3"),
        ("2. Visión general del proyecto", "4"),
        ("3. Stack técnico y decisiones de arquitectura", "5"),
        ("4. Estructura del proyecto", "6"),
        ("5. Modelo de datos (con diagrama)", "7"),
        ("6. Exposición por integrante", "9"),
        ("    6.1 Airton Rodriguez — Base de Datos", "9"),
        ("    6.2 Carlos Huane — Arquitectura y Seguridad JWT", "12"),
        ("    6.3 Bartolome Prado — Tickets, SLA, Notificaciones", "16"),
        ("    6.4 Matias Rodriguez — Usuarios, Categorías, KB", "20"),
        ("    6.5 Claudia Gonzales — Dashboard, Ranking, Export", "23"),
        ("7. Anexo A — Cómo levantar el proyecto", "26"),
        ("8. Anexo B — Credenciales de prueba", "27"),
        ("9. Anexo C — Listado completo de endpoints", "28"),
    ]
    for item, pag in indice:
        p = doc.add_paragraph()
        r = p.add_run(f"{item}")
        r.font.size = Pt(11)
        r2 = p.add_run(f"  .....  pág. {pag}")
        r2.font.size = Pt(10)
        r2.font.color.rgb = COLOR_GRIS

    saltopagina(doc)

    # ═══════════════════════════════════════════════════════════════
    # 1. CÓMO USAR ESTA GUÍA
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "1. Cómo usar esta guía", nivel=1)
    parrafo(doc,
            "Este documento está pensado para que cada integrante del Grupo 6 prepare su parte de la "
            "exposición del backend. La estructura es:")
    bullet(doc, "Las secciones 2-5 son CONOCIMIENTO COMPARTIDO — todos deben dominarlas para responder "
                 "preguntas del docente sobre el proyecto en general.")
    bullet(doc, "La sección 6 está dividida en 5 sub-secciones, una por integrante. Cada quien estudia "
                 "su sub-sección a fondo: archivos a abrir, código a explicar, demo en Swagger.")
    bullet(doc, "Los anexos (A, B, C) son referencia rápida durante la exposición.")

    parrafo(doc, "Para cada integrante encontrarás:", size=11, bold=True)
    bullet(doc, "Archivos a abrir en VS Code (rutas exactas).")
    bullet(doc, "Código clave explicado línea por línea.")
    bullet(doc, "Endpoints a demostrar en Swagger UI con ejemplos de request/response.")
    bullet(doc, "Preguntas frecuentes del docente con respuestas sugeridas.")

    callout(doc, "Consejo de exposición",
            "Antes de exponer, todos deben haber levantado el proyecto al menos una vez y haber "
            "ejecutado 1-2 endpoints de su módulo en Swagger. La demostración en vivo siempre vale "
            "más que las diapositivas.")

    saltopagina(doc)

    # ═══════════════════════════════════════════════════════════════
    # 2. VISIÓN GENERAL
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "2. Visión general del proyecto", nivel=1)
    parrafo(doc,
            "TelecoPerú S.A.C. es una empresa de telecomunicaciones que necesita una Intranet para gestionar "
            "incidencias internas (tickets de soporte). Los empleados pueden reportar problemas técnicos, "
            "los técnicos los atienden y los administradores supervisan métricas de cumplimiento (SLA).")

    parrafo(doc, "Objetivos del sistema:", size=11, bold=True)
    bullet(doc, "Permitir a clientes (empleados) crear y dar seguimiento a sus tickets.")
    bullet(doc, "Asignar tickets a técnicos según categoría y prioridad.")
    bullet(doc, "Monitorear el cumplimiento del SLA (Service Level Agreement) por prioridad.")
    bullet(doc, "Mantener una base de conocimiento (artículos de autoservicio) para reducir tickets.")
    bullet(doc, "Generar reportes y estadísticas para la gerencia de TI.")

    parrafo(doc, "Roles del sistema:", size=11, bold=True)
    tabla_simple(doc,
                 ["Rol", "Qué puede hacer"],
                 [
                     ["ADMIN", "Acceso total: gestión de usuarios, categorías, SLA, ver todo, exportar, dashboard."],
                     ["TECNICO", "Atender tickets asignados, cambiar estado, ver dashboard, escribir artículos KB."],
                     ["CLIENTE", "Crear tickets, ver sus propios tickets, leer la base de conocimiento."],
                 ],
                 anchos=[3, 13])

    saltopagina(doc)

    # ═══════════════════════════════════════════════════════════════
    # 3. STACK TÉCNICO
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "3. Stack técnico y decisiones de arquitectura", nivel=1)

    tabla_simple(doc,
                 ["Capa", "Tecnología", "Versión"],
                 [
                     ["Lenguaje", "Java", "21"],
                     ["Framework", "Spring Boot", "3.3.5"],
                     ["Web", "Spring Web (REST)", "—"],
                     ["Persistencia", "Spring Data JPA + Hibernate", "6.5.3"],
                     ["Base de datos", "MySQL", "8.0+"],
                     ["Seguridad", "Spring Security + JWT (jjwt)", "0.11.5"],
                     ["Validaciones", "Jakarta Validation (Hibernate Validator)", "—"],
                     ["Build", "Maven", "3.9+"],
                     ["Documentación API", "SpringDoc OpenAPI (Swagger UI)", "2.5.0"],
                     ["Correo", "Spring Mail (JavaMailSender)", "—"],
                     ["Excel", "Apache POI", "5.2.5"],
                     ["PDF", "OpenPDF", "1.3.43"],
                 ],
                 anchos=[3.5, 8, 3])

    heading(doc, "¿Por qué API REST y no MVC con vistas?", nivel=2)
    parrafo(doc,
            "El frontend de la Intranet será desarrollado en Angular (proyecto independiente). El backend solo "
            "expone una API REST que devuelve JSON. Por eso usamos @RestController y no @Controller con "
            "vistas Thymeleaf. Las páginas web las renderiza Angular en el navegador del cliente, mientras "
            "que Spring Boot solo entrega datos.")

    heading(doc, "¿Por qué JWT y no sesiones?", nivel=2)
    parrafo(doc,
            "Las SPA (Single Page Applications) como Angular se llevan mal con cookies de sesión por temas "
            "de CORS y escalabilidad horizontal. JWT permite que el cliente envíe el token en cada petición "
            "(en el header Authorization), lo que vuelve al servidor stateless: cualquier instancia del "
            "backend puede atender la petición sin compartir estado de sesión.")

    callout(doc, "Punto clave para la exposición",
            "Si el docente pregunta por qué no usamos formLogin() como en el ejemplo de clase: "
            "responder que el ejemplo Mystore es Spring MVC server-side con Thymeleaf, mientras "
            "que nuestro proyecto es API REST puro para Angular. Ambos patrones son válidos, pero "
            "el nuestro corresponde al stack real de la asignatura (frontend separado).")

    saltopagina(doc)

    # ═══════════════════════════════════════════════════════════════
    # 4. ESTRUCTURA DEL PROYECTO
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "4. Estructura del proyecto", nivel=1)

    parrafo(doc, "Patrón de capas:", size=11, bold=True)
    codigo(doc,
"""Cliente HTTP (Postman / Angular / Swagger)
         |
         v
  [ Controller ]   <- recibe request, devuelve JSON
         |
         v
   [ Service ]     <- logica de negocio + @Transactional
         |
         v
  [ Repository ]   <- acceso a datos (JPA)
         |
         v
    [ MySQL ]      <- intranet_db""")

    parrafo(doc, "Estructura de carpetas:", size=11, bold=True)
    codigo(doc,
"""src/main/java/com/grupo6/intranet/
├── config/
│   ├── DataInitializer.java        # Datos de prueba al primer arranque
│   ├── GlobalExceptionHandler.java # Manejo centralizado de errores
│   ├── JwtFilter.java              # Intercepta cada request y valida JWT
│   ├── JwtUtil.java                # Genera y parsea tokens JWT
│   ├── OpenApiConfig / SwaggerConfig.java # Configuracion Swagger
│   ├── SecurityConfig.java         # Reglas de autorizacion por rol
│   └── SlaScheduler.java           # Tarea programada de alertas SLA
├── controllers/
│   ├── AuthController.java         # /api/auth/login y /api/auth/registro
│   ├── ArticuloController.java     # /api/articulos (KB)
│   ├── CategoriaController.java    # /api/categorias
│   ├── DashboardController.java    # /api/dashboard
│   ├── ExportController.java       # /api/export
│   ├── RankingController.java      # /api/ranking
│   ├── SlaConfigController.java    # /api/sla
│   ├── TicketController.java       # /api/tickets
│   └── UsuarioController.java      # /api/usuarios
├── dtos/                           # Objetos de transferencia (Request/Response)
├── models/                         # Entidades JPA (@Entity)
├── repositories/                   # Interfaces JpaRepository
├── services/                       # Logica de negocio
└── IntranetApplication.java        # main()""")

    saltopagina(doc)

    # ═══════════════════════════════════════════════════════════════
    # 5. MODELO DE DATOS
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "5. Modelo de datos (intranet_db)", nivel=1)
    parrafo(doc,
            "La base de datos consta de 7 tablas relacionadas. El esquema se genera automáticamente "
            "al arrancar la aplicación gracias a Hibernate (spring.jpa.hibernate.ddl-auto=update) y los "
            "datos iniciales son sembrados por DataInitializer.java.")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture("docs/diagrama-bd.png", width=Cm(16))
    parrafo(doc, "Figura 1. Diagrama Entidad-Relación de intranet_db", size=9,
            italic=True, color=COLOR_GRIS, align=WD_ALIGN_PARAGRAPH.CENTER)

    heading(doc, "Resumen de tablas", nivel=2)
    tabla_simple(doc,
                 ["Tabla", "Descripción", "Registros iniciales"],
                 [
                     ["usuarios", "Empleados de TelecoPerú (admins, técnicos, clientes).", "7"],
                     ["categorias", "Categorías de incidencias (Red, Hardware, Software, etc.).", "6"],
                     ["subcategorias", "Subdivisión de cada categoría.", "12"],
                     ["tickets", "Incidencias creadas por clientes.", "6"],
                     ["historial_tickets", "Cambios de estado de cada ticket.", "13"],
                     ["sla_config", "Tiempos SLA por prioridad.", "5"],
                     ["articulos", "Base de conocimiento (autoservicio).", "5"],
                 ],
                 anchos=[3.5, 9, 3])

    heading(doc, "Relaciones clave", nivel=2)
    bullet(doc, "Un usuario puede ser CLIENTE de muchos tickets (tickets.cliente_id).")
    bullet(doc, "Un usuario puede ser TÉCNICO asignado a muchos tickets (tickets.tecnico_id).")
    bullet(doc, "Una categoría tiene muchas subcategorías (subcategorias.categoria_id).")
    bullet(doc, "Un ticket tiene un historial de cambios (historial_tickets.ticket_id).")
    bullet(doc, "Cada artículo de KB pertenece a una categoría y tiene un autor (usuario).")

    saltopagina(doc)

    # ═══════════════════════════════════════════════════════════════
    # 6. EXPOSICIÓN POR INTEGRANTE
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "6. Exposición por integrante", nivel=1)

    parrafo(doc,
            "Cada integrante cuenta una historia completa: del modelo de datos al endpoint funcionando "
            "en Swagger. Esto facilita que el docente entienda el módulo sin tener que conectar piezas.")

    parrafo(doc, "División acordada:", size=11, bold=True)
    tabla_simple(doc,
                 ["#", "Integrante", "Módulo a exponer"],
                 [
                     ["6.1", "Airton Rodriguez", "Modelo de Datos + DataInitializer"],
                     ["6.2", "Carlos Huane", "Arquitectura + Seguridad JWT"],
                     ["6.3", "Bartolome Prado", "Tickets + SLA + Scheduler + Email"],
                     ["6.4", "Matias Rodriguez", "Usuarios + Categorías + Base de Conocimiento"],
                     ["6.5", "Claudia Gonzales", "Dashboard + Ranking + Exportación"],
                 ],
                 anchos=[1, 5, 9])

    saltopagina(doc)

    # ───────────────────────────────────────────────────────────────
    # 6.1 AIRTON — BASE DE DATOS
    # ───────────────────────────────────────────────────────────────
    heading(doc, "6.1 Airton Rodriguez — Modelo de Datos y DataInitializer", nivel=2)
    parrafo(doc, "Duración sugerida: 5 minutos", italic=True, color=COLOR_GRIS)

    heading(doc, "Qué expones", nivel=3)
    bullet(doc, "Estructura física de la BD intranet_db en MySQL.")
    bullet(doc, "Las 7 tablas, sus columnas, claves primarias y foráneas.")
    bullet(doc, "Cómo Hibernate genera el esquema automáticamente al arrancar.")
    bullet(doc, "Cómo DataInitializer.java siembra registros de prueba.")

    heading(doc, "Archivos a abrir en VS Code", nivel=3)
    codigo(doc,
"""src/main/java/com/grupo6/intranet/models/
├── Usuario.java       <- Mostrar @Entity, @Table, @Column, @Email, validaciones
├── Categoria.java     <- Tabla simple
├── Subcategoria.java  <- @ManyToOne hacia Categoria
├── Ticket.java        <- @ManyToOne multiples (cliente, tecnico, categoria)
├── HistorialTicket.java
├── SlaConfig.java
├── Articulo.java
├── Rol.java           <- ENUM: ADMIN, TECNICO, CLIENTE
├── Estado.java        <- ENUM: NUEVO, EN_ATENCION, ESCALADO, RESUELTO, CERRADO
└── Prioridad.java     <- ENUM: CRITICA, ALTA, MEDIA, BAJA, SIN_ASIGNAR

src/main/java/com/grupo6/intranet/config/DataInitializer.java""")

    heading(doc, "Código clave a explicar — anatomía de una entidad JPA", nivel=3)
    codigo(doc,
"""@Data                                       // Lombok: getters, setters, toString
@Entity                                     // JPA: esta clase mapea a una tabla
@Table(name = "tickets")                    // Nombre de la tabla en MySQL
public class Ticket {

    @Id                                     // Clave primaria
    @GeneratedValue(strategy = GenerationType.IDENTITY)  // AUTO_INCREMENT
    private Long id;

    @NotEmpty(message = "El título es obligatorio")     // Validacion Jakarta
    @Column(nullable = false)
    private String titulo;

    @Column(nullable = false, columnDefinition = "TEXT")
    private String descripcion;

    @ManyToOne                              // Relacion N:1
    @JoinColumn(name = "cliente_id", nullable = false)
    private Usuario cliente;                // FK a usuarios

    @Enumerated(EnumType.STRING)            // Guarda "NUEVO" en vez de 0
    private Estado estado = Estado.NUEVO;

    @PrePersist                             // Hook antes de INSERT
    protected void onCreate() {
        createdAt = LocalDateTime.now();
    }
}""")

    heading(doc, "Código clave a explicar — DataInitializer", nivel=3)
    codigo(doc,
"""@Component                                  // Spring lo detecta y ejecuta
public class DataInitializer implements CommandLineRunner {

    @Override
    public void run(String... args) {
        if (usuarioRepo.count() > 0) return; // Idempotente: solo siembra 1 vez

        Usuario admin = usuario("Carlos", "Ramirez",
                "admin@telecoperu.com", "admin123", "987654321", Rol.ADMIN);
        // ... mas usuarios, categorias, tickets, articulos
    }

    private Usuario usuario(String nombre, String apellido, ...) {
        Usuario u = new Usuario();
        u.setPassword(passwordEncoder.encode(pass));  // BCrypt
        return usuarioRepo.save(u);
    }
}""")

    heading(doc, "Demo en DBeaver / MySQL Workbench", nivel=3)
    bullet(doc, "Conectarse a MySQL local (host: localhost, port: 3306, user: root).")
    bullet(doc, "Mostrar la BD intranet_db y sus 7 tablas en el explorador.")
    bullet(doc, 'Ejecutar: SELECT * FROM usuarios; — mostrar los 7 registros sembrados.')
    bullet(doc, 'Ejecutar: SELECT * FROM tickets; — mostrar los 6 tickets de prueba.')
    bullet(doc, 'Ejecutar: SELECT t.titulo, u.nombre FROM tickets t JOIN usuarios u ON u.id = t.cliente_id; '
                 "— mostrar JOIN.")

    heading(doc, "Posibles preguntas del docente", nivel=3)
    tabla_simple(doc,
                 ["Pregunta", "Respuesta sugerida"],
                 [
                     ["¿Quién crea las tablas?",
                      "Hibernate, gracias a spring.jpa.hibernate.ddl-auto=update. Lee las clases @Entity y "
                      "ejecuta los CREATE TABLE / ALTER TABLE necesarios."],
                     ["¿Por qué usan ENUM en Estado y Prioridad?",
                      "Porque son listas cerradas de valores. @Enumerated(EnumType.STRING) guarda el nombre "
                      "(NUEVO) en lugar del índice (0), lo que hace la BD más legible y resistente a cambios "
                      "en el orden del enum."],
                     ["¿Cómo se garantiza que admin@telecoperu.com es único?",
                      "Con @Column(unique = true) en el campo email de Usuario. Hibernate crea un UNIQUE INDEX."],
                     ["¿Por qué la password no aparece en el JSON cuando consultan usuarios?",
                      "Tiene @JsonIgnore en la entidad Usuario, así Jackson la excluye al serializar."],
                 ],
                 anchos=[5, 11])

    saltopagina(doc)

    # ───────────────────────────────────────────────────────────────
    # 6.2 CARLOS — ARQUITECTURA + SEGURIDAD
    # ───────────────────────────────────────────────────────────────
    heading(doc, "6.2 Carlos Huane — Arquitectura y Seguridad JWT", nivel=2)
    parrafo(doc, "Duración sugerida: 7 minutos", italic=True, color=COLOR_GRIS)

    heading(doc, "Qué expones", nivel=3)
    bullet(doc, "Visión global del proyecto: pom.xml, estructura de carpetas, patrón de capas.")
    bullet(doc, "Cómo funciona la autenticación con JWT (login → token → header en cada request).")
    bullet(doc, "Cómo Spring Security autoriza los endpoints según el rol del usuario.")
    bullet(doc, "El endpoint POST /api/auth/registro con BindingResult (estilo del docente).")

    heading(doc, "Archivos a abrir en VS Code", nivel=3)
    codigo(doc,
"""pom.xml                                                  <- Dependencias Maven
src/main/resources/application.properties                <- Config DB + JWT + Mail
src/main/java/com/grupo6/intranet/
├── IntranetApplication.java                             <- main + @EnableScheduling
└── config/
    ├── SecurityConfig.java        <- Reglas de autorizacion (hasRole)
    ├── JwtUtil.java               <- Genera y valida tokens
    ├── JwtFilter.java             <- Intercepta cada request
    ├── GlobalExceptionHandler.java <- Errores centralizados
    └── SwaggerConfig.java
└── controllers/AuthController.java  <- login + registro""")

    heading(doc, "Flujo de autenticación (explicar con diagrama dibujado a mano)", nivel=3)
    codigo(doc,
"""1) POST /api/auth/login  { email, password }
   |
   v
2) UsuarioService busca por email + BCrypt.matches(password)
   |
   v
3) Si OK -> JwtUtil.generarToken(email, rol)
   |
   v
4) Cliente guarda token y lo envia en cada request:
   Authorization: Bearer eyJhbGc...
   |
   v
5) JwtFilter intercepta, valida el token, mete el rol en SecurityContext
   |
   v
6) SecurityConfig revisa si el rol tiene permiso para la ruta
   |
   v
7) Controller atiende la peticion""")

    heading(doc, "Código clave — JwtUtil", nivel=3)
    codigo(doc,
"""public String generarToken(String email, String rol) {
    return Jwts.builder()
            .setSubject(email)                                  // "sub"
            .claim("rol", rol)                                  // claim custom
            .setIssuedAt(new Date())                            // "iat"
            .setExpiration(new Date(System.currentTimeMillis() + expiration))
            .signWith(getKey(), SignatureAlgorithm.HS256)       // Firma HMAC-SHA256
            .compact();
}""")

    heading(doc, "Código clave — JwtFilter", nivel=3)
    codigo(doc,
"""@Override
protected void doFilterInternal(HttpServletRequest request, ...) {
    String header = request.getHeader("Authorization");
    if (header != null && header.startsWith("Bearer ")) {
        String token = header.substring(7);
        if (jwtUtil.esValido(token)) {
            String email = jwtUtil.extraerEmail(token);
            String rol = jwtUtil.extraerRol(token);
            UsernamePasswordAuthenticationToken auth = new UsernamePasswordAuthenticationToken(
                email, null, List.of(new SimpleGrantedAuthority("ROLE_" + rol))
            );
            SecurityContextHolder.getContext().setAuthentication(auth);
        }
    }
    chain.doFilter(request, response);
}""")

    heading(doc, "Código clave — SecurityConfig con roles", nivel=3)
    codigo(doc,
""".authorizeHttpRequests(auth -> auth
    .requestMatchers("/api/auth/**", "/swagger-ui/**", "/api-docs/**").permitAll()
    .requestMatchers(HttpMethod.GET, "/api/categorias/**").permitAll()

    .requestMatchers(HttpMethod.GET, "/api/usuarios/**").hasAnyRole("ADMIN", "TECNICO")
    .requestMatchers("/api/usuarios/**").hasRole("ADMIN")

    .requestMatchers("/api/dashboard/**").hasAnyRole("ADMIN", "TECNICO")
    .requestMatchers("/api/sla/**").hasRole("ADMIN")

    .anyRequest().authenticated()
)""")

    heading(doc, "Código clave — registro con BindingResult (estilo del docente)", nivel=3)
    codigo(doc,
"""@PostMapping("/registro")
public ResponseEntity<?> registro(@Valid @RequestBody RegistroClienteRequest req,
                                  BindingResult result) {
    if (!req.getPassword().equals(req.getConfirmPassword())) {
        result.addError(new FieldError(
            "registroClienteRequest", "confirmPassword",
            "Las contraseñas no coinciden"));
    }
    if (usuarioService.buscarPorEmail(req.getEmail()).isPresent()) {
        result.addError(new FieldError(
            "registroClienteRequest", "email",
            "El email ya está registrado"));
    }
    if (result.hasErrors()) {
        // construye Map<campo, mensaje> y devuelve 400
    }
    Usuario nuevo = ...; // crear con rol CLIENTE
    String token = jwtUtil.generarToken(...);
    return ResponseEntity.ok(new LoginResponse(token, ...));
}""")

    heading(doc, "Demo en Swagger UI", nivel=3)
    bullet(doc, "Abrir http://localhost:8080/swagger-ui.html")
    bullet(doc, "POST /api/auth/registro con body: { nombre, apellido, email, password, confirmPassword, telefono }")
    bullet(doc, "Mostrar caso de éxito y caso de error (contraseñas no coinciden) — el JSON de error es legible.")
    bullet(doc, "POST /api/auth/login con admin@telecoperu.com / admin123 → copiar token.")
    bullet(doc, 'Click en "Authorize 🔓" arriba a la derecha, pegar token, Authorize.')
    bullet(doc, "GET /api/usuarios → 200 OK con la lista (porque admin tiene permiso).")
    bullet(doc, "Logout en Swagger, login con cliente@... → GET /api/usuarios → 403 Forbidden.")

    heading(doc, "Posibles preguntas del docente", nivel=3)
    tabla_simple(doc,
                 ["Pregunta", "Respuesta sugerida"],
                 [
                     ["¿Qué es JWT?",
                      "JSON Web Token: un token firmado que contiene claims (email, rol). El servidor no "
                      "guarda sesión, todo el contexto va en el token. Es el estándar de la industria para "
                      "SPAs y APIs móviles."],
                     ["¿Cómo se firma el token?",
                      "Con HMAC-SHA256 usando la clave secreta jwt.secret definida en application.properties. "
                      "El servidor puede verificar la firma sin contactar a una BD de sesiones."],
                     ["¿Por qué el token expira en 24 horas?",
                      "jwt.expiration=86400000 ms. Es un balance entre seguridad (token corto) y comodidad "
                      "(no obligar al usuario a re-loguearse cada hora). En producción se complementa con "
                      "refresh tokens."],
                     ["¿Por qué hay @Transactional en los servicios?",
                      "Para garantizar atomicidad: si crear un ticket también requiere insertar en "
                      "historial_tickets, ambas operaciones se confirman o se revierten juntas. "
                      "@Transactional(readOnly=true) a nivel clase optimiza lecturas; los métodos de "
                      "escritura lo sobreescriben."],
                     ["¿Por qué BindingResult solo en /api/auth/registro?",
                      "En el resto usamos @Valid sin BindingResult: cuando hay error, Spring lanza "
                      "MethodArgumentNotValidException y GlobalExceptionHandler la traduce a un JSON. "
                      "En /registro usamos BindingResult porque necesitamos validar dos campos cruzados "
                      "(password vs confirmPassword) que @Valid solo no detecta."],
                 ],
                 anchos=[5, 11])

    saltopagina(doc)

    # ───────────────────────────────────────────────────────────────
    # 6.3 BARTOLOME — TICKETS + SLA + EMAIL
    # ───────────────────────────────────────────────────────────────
    heading(doc, "6.3 Bartolome Prado — Tickets, SLA, Notificaciones", nivel=2)
    parrafo(doc, "Duración sugerida: 7 minutos", italic=True, color=COLOR_GRIS)

    heading(doc, "Qué expones", nivel=3)
    bullet(doc, "El ciclo de vida de un ticket: NUEVO → EN_ATENCION → RESUELTO → CERRADO.")
    bullet(doc, "Cómo se registra el historial automáticamente en cada cambio de estado.")
    bullet(doc, "Configuración SLA: tiempos por prioridad (CRITICA, ALTA, MEDIA, BAJA).")
    bullet(doc, "Scheduler que cada 30 minutos revisa tickets a punto de vencer SLA.")
    bullet(doc, "Notificaciones por correo en eventos clave (creación, asignación, cambio de estado).")

    heading(doc, "Archivos a abrir en VS Code", nivel=3)
    codigo(doc,
"""src/main/java/com/grupo6/intranet/
├── models/
│   ├── Ticket.java
│   ├── HistorialTicket.java
│   └── SlaConfig.java
├── repositories/
│   ├── TicketRepository.java           <- queries JPQL y nativas
│   └── SlaConfigRepository.java
├── services/
│   ├── TicketService.java              <- crear, cambiarEstado
│   └── EmailService.java               <- plantillas HTML
├── controllers/
│   ├── TicketController.java
│   └── SlaConfigController.java
└── config/
    └── SlaScheduler.java               <- @Scheduled cada 30min""")

    heading(doc, "Código clave — TicketService.crear con transacción", nivel=3)
    codigo(doc,
"""@Transactional                              // sesion 7 del silabo
public Ticket crear(TicketRequest req, Long clienteId) {
    Usuario cliente = usuarioRepository.findById(clienteId)
            .orElseThrow(() -> new RuntimeException("Cliente no encontrado"));
    Categoria categoria = categoriaRepository.findById(req.getCategoriaId())
            .orElseThrow(() -> new RuntimeException("Categoría no encontrada"));

    Ticket ticket = new Ticket();
    ticket.setTitulo(req.getTitulo());
    ticket.setDescripcion(req.getDescripcion());
    ticket.setCliente(cliente);
    ticket.setCategoria(categoria);
    ticket.setPrioridad(req.getPrioridad() != null ? req.getPrioridad() : Prioridad.SIN_ASIGNAR);
    ticket.setEstado(Estado.NUEVO);

    Ticket guardado = ticketRepository.save(ticket);
    registrarHistorial(guardado, cliente, null, Estado.NUEVO, "Ticket creado");
    emailService.notificarTicketCreado(guardado);           // dispara correo
    return guardado;
}""")

    heading(doc, "Código clave — Scheduler SLA", nivel=3)
    codigo(doc,
"""@Scheduled(cron = "0 0/30 * * * *")          // Cada 30 min
public void revisarSlaTickets() {
    List<Ticket> tickets = ticketRepository.findAbiertosSinAlerta();
    LocalDateTime ahora = LocalDateTime.now();

    for (Ticket ticket : tickets) {
        SlaConfig sla = slaConfigRepository.findByPrioridad(ticket.getPrioridad()).get();
        int horasResolucion = sla.getTiempoResolucionHoras();
        long horasTranscurridas = Duration.between(ticket.getCreatedAt(), ahora).toHours();
        long horasRestantes = horasResolucion - horasTranscurridas;

        if (horasRestantes <= horasResolucion * 0.25 && horasRestantes > 0) {
            emailService.notificarAlertaSla(ticket, (int) horasRestantes);
            ticket.setAlertaSlaEnviada(true);                // evita reenviar
            ticketRepository.save(ticket);
        }
    }
}""")

    heading(doc, "Demo en Swagger UI", nivel=3)
    bullet(doc, "Login como CLIENTE: pedro.sanchez@telecoperu.com / cliente123. Copiar token.")
    bullet(doc, "POST /api/tickets/cliente/5 (id de Pedro) con body de ticket nuevo.")
    bullet(doc, "Logout. Login como TECNICO: tecnico1@telecoperu.com / tecnico123.")
    bullet(doc, "PATCH /api/tickets/{ticketId}/estado/3 (id del tecnico) con estadoNuevo=EN_ATENCION + tecnicoId=3 + comentario.")
    bullet(doc, "GET /api/tickets/{id}/historial — mostrar los registros.")
    bullet(doc, "GET /api/sla — mostrar la matriz de tiempos.")

    heading(doc, "Posibles preguntas del docente", nivel=3)
    tabla_simple(doc,
                 ["Pregunta", "Respuesta sugerida"],
                 [
                     ["¿Qué pasa si crean un ticket pero falla la inserción del historial?",
                      "Por @Transactional, toda la operación se revierte (rollback). Ni el ticket ni el "
                      "historial quedan en la BD. Esto evita inconsistencias."],
                     ["¿Por qué el correo no rompe la creación del ticket si falla SMTP?",
                      "EmailService captura las excepciones internamente y solo las loggea. El método que "
                      "lo llama no recibe error, así que la transacción del ticket sigue su curso normal."],
                     ["¿Cómo se asegura el scheduler de no enviar el mismo correo varias veces?",
                      "El campo alertaSlaEnviada (Boolean) en la tabla tickets se pone en true cuando se "
                      "envía la alerta. La query findAbiertosSinAlerta() filtra por ese campo."],
                     ["¿Qué pasa si el servidor se reinicia con tickets vencidos?",
                      "El scheduler corre al iniciar el proceso (cron 0 0/30) o en la próxima media hora. "
                      "Detecta los tickets sin alerta y los procesa."],
                 ],
                 anchos=[5, 11])

    saltopagina(doc)

    # ───────────────────────────────────────────────────────────────
    # 6.4 MATIAS — USUARIOS + CATEGORÍAS + KB
    # ───────────────────────────────────────────────────────────────
    heading(doc, "6.4 Matias Rodriguez — Usuarios, Categorías, Base de Conocimiento", nivel=2)
    parrafo(doc, "Duración sugerida: 6 minutos", italic=True, color=COLOR_GRIS)

    heading(doc, "Qué expones", nivel=3)
    bullet(doc, "CRUD de usuarios (alta, baja, activar/desactivar).")
    bullet(doc, "Catálogo de categorías y subcategorías para clasificar tickets.")
    bullet(doc, "Base de conocimiento: artículos de autoservicio que clientes pueden consultar antes de "
                 "abrir un ticket. Incluye contador de vistas y búsqueda por texto.")

    heading(doc, "Archivos a abrir en VS Code", nivel=3)
    codigo(doc,
"""src/main/java/com/grupo6/intranet/
├── models/
│   ├── Usuario.java
│   ├── Categoria.java
│   ├── Subcategoria.java
│   └── Articulo.java
├── repositories/
│   ├── UsuarioRepository.java
│   ├── CategoriaRepository.java
│   └── ArticuloRepository.java         <- @Query busqueda por texto
├── services/
│   ├── UsuarioService.java
│   ├── CategoriaService.java
│   └── ArticuloService.java            <- buscarPorId incrementa vistas
├── controllers/
│   ├── UsuarioController.java
│   ├── CategoriaController.java
│   └── ArticuloController.java
└── dtos/
    └── ArticuloRequest.java""")

    heading(doc, "Código clave — Repositorio con @Query (sesión 7 del sílabo: JPQL)", nivel=3)
    codigo(doc,
"""public interface ArticuloRepository extends JpaRepository<Articulo, Long> {

    List<Articulo> findByActivoTrueOrderByCreatedAtDesc();

    List<Articulo> findByCategoriaIdAndActivoTrueOrderByCreatedAtDesc(Long categoriaId);

    @Query("SELECT a FROM Articulo a WHERE a.activo = true AND " +     // JPQL
           "(LOWER(a.titulo) LIKE LOWER(CONCAT('%', :texto, '%')) OR " +
           "LOWER(a.contenido) LIKE LOWER(CONCAT('%', :texto, '%'))) " +
           "ORDER BY a.vistas DESC")
    List<Articulo> buscarPorTexto(@Param("texto") String texto);

    List<Articulo> findTop10ByActivoTrueOrderByVistasDesc();
}""")

    heading(doc, "Código clave — Service con @Transactional", nivel=3)
    codigo(doc,
"""@Service
@Transactional(readOnly = true)              // por defecto, lecturas
public class ArticuloService {

    @Transactional                            // sobreescribe para escritura
    public Optional<Articulo> buscarPorId(Long id) {
        return articuloRepository.findById(id).map(a -> {
            a.setVistas(a.getVistas() + 1);   // efecto colateral: incrementa
            return articuloRepository.save(a);
        });
    }
}""")

    heading(doc, "Demo en Swagger UI", nivel=3)
    bullet(doc, "POST /api/auth/registro con un cliente nuevo → muestra creación pública.")
    bullet(doc, "GET /api/categorias (sin token) → lista categorías activas (es público).")
    bullet(doc, "Login como ADMIN, GET /api/usuarios → lista los 7+ usuarios.")
    bullet(doc, "GET /api/articulos → lista los 5 artículos sembrados.")
    bullet(doc, 'GET /api/articulos/buscar?q=vpn → busca por texto y devuelve el que tiene "VPN" en el título.')
    bullet(doc, "GET /api/articulos/3 dos veces → notar que vistas aumenta.")
    bullet(doc, "POST /api/articulos/autor/3 (técnico) con un artículo nuevo.")

    heading(doc, "Posibles preguntas del docente", nivel=3)
    tabla_simple(doc,
                 ["Pregunta", "Respuesta sugerida"],
                 [
                     ["¿Qué es JPQL y en qué se diferencia de SQL?",
                      "JPQL (Java Persistence Query Language) trabaja con entidades en vez de tablas. "
                      "En lugar de SELECT * FROM articulos escribimos SELECT a FROM Articulo a. "
                      "Hibernate lo traduce a SQL nativo de MySQL en tiempo de ejecución."],
                     ["¿Para qué sirven los métodos derivados como findByActivoTrue?",
                      "Spring Data JPA genera el SQL automáticamente a partir del nombre del método. "
                      "Es útil para consultas simples sin tener que escribir JPQL."],
                     ["¿Por qué buscarPorId incrementa vistas si solo es un GET?",
                      "Decisión de diseño: cada lectura de un artículo cuenta como una visita. Por eso el "
                      "método está marcado con @Transactional (no readOnly) para permitir el UPDATE."],
                     ["¿Cómo se valida que el email del usuario es único?",
                      "@Column(unique = true) en la entidad + manejo de la excepción en el service. "
                      "En el endpoint /api/auth/registro lo validamos manualmente antes para devolver un "
                      "mensaje amigable."],
                 ],
                 anchos=[5, 11])

    saltopagina(doc)

    # ───────────────────────────────────────────────────────────────
    # 6.5 CLAUDIA — DASHBOARD + RANKING + EXPORT
    # ───────────────────────────────────────────────────────────────
    heading(doc, "6.5 Claudia Gonzales — Dashboard, Ranking y Exportación", nivel=2)
    parrafo(doc, "Duración sugerida: 6 minutos", italic=True, color=COLOR_GRIS)

    heading(doc, "Qué expones", nivel=3)
    bullet(doc, "Endpoints de estadísticas para que el frontend Angular construya gráficos.")
    bullet(doc, "Consultas JPQL agregadas (GROUP BY, COUNT, AVG).")
    bullet(doc, "Ranking de categorías más reportadas y técnicos más productivos.")
    bullet(doc, "Generación de reportes Excel (Apache POI) y PDF (OpenPDF) descargables.")

    heading(doc, "Archivos a abrir en VS Code", nivel=3)
    codigo(doc,
"""src/main/java/com/grupo6/intranet/
├── repositories/
│   └── TicketRepository.java     <- queries con COUNT, AVG, GROUP BY
├── services/
│   ├── DashboardService.java
│   ├── RankingService.java
│   └── ExportService.java        <- Apache POI + OpenPDF
├── controllers/
│   ├── DashboardController.java
│   ├── RankingController.java
│   └── ExportController.java
└── dtos/
    ├── ResumenDashboardDto.java
    ├── ConteoDto.java
    ├── TecnicoConteoDto.java
    ├── CategoriaRankingDto.java
    └── TecnicoRankingDto.java""")

    heading(doc, "Código clave — Query agregada JPQL", nivel=3)
    codigo(doc,
"""@Query("SELECT t.estado, COUNT(t) FROM Ticket t GROUP BY t.estado")
List<Object[]> contarPorEstado();

@Query("SELECT t.tecnico.id, t.tecnico.nombre, t.tecnico.apellido, COUNT(t) " +
       "FROM Ticket t WHERE t.tecnico IS NOT NULL " +
       "GROUP BY t.tecnico.id, t.tecnico.nombre, t.tecnico.apellido " +
       "ORDER BY COUNT(t) DESC")
List<Object[]> contarPorTecnico();""")

    heading(doc, "Código clave — Query nativa para AVG con TIMESTAMPDIFF", nivel=3)
    parrafo(doc,
            "Hibernate 6 no infiere el tipo de retorno de FUNCTION() dentro de AVG(). "
            "Por eso usamos nativeQuery=true en las consultas que calculan promedios de tiempo.",
            italic=True, color=COLOR_GRIS)
    codigo(doc,
"""@Query(value = "SELECT c.id, c.nombre, COUNT(t.id), " +
       "AVG(TIMESTAMPDIFF(HOUR, t.created_at, t.fecha_resolucion)) " +
       "FROM tickets t JOIN categorias c ON c.id = t.categoria_id " +
       "WHERE t.fecha_resolucion IS NOT NULL " +
       "GROUP BY c.id, c.nombre " +
       "ORDER BY COUNT(t.id) DESC", nativeQuery = true)
List<Object[]> rankingCategoriasConTiempo();""")

    heading(doc, "Código clave — Service con DTO", nivel=3)
    codigo(doc,
"""public List<TecnicoConteoDto> ticketsPorTecnico() {
    return ticketRepository.contarPorTecnico().stream()
            .map(row -> new TecnicoConteoDto(
                    (Long) row[0],
                    row[1] + " " + row[2],
                    (Long) row[3]))
            .collect(Collectors.toList());
}""")

    heading(doc, "Código clave — Exportar Excel con Apache POI", nivel=3)
    codigo(doc,
"""public byte[] exportarTicketsExcel(Estado estado, LocalDateTime desde, LocalDateTime hasta)
        throws IOException {
    List<Ticket> tickets = ticketRepository.filtrar(estado, desde, hasta);
    try (Workbook wb = new XSSFWorkbook();
         ByteArrayOutputStream out = new ByteArrayOutputStream()) {

        Sheet sheet = wb.createSheet("Tickets");
        Row header = sheet.createRow(0);
        header.createCell(0).setCellValue("ID");
        header.createCell(1).setCellValue("Título");
        // ... mas columnas

        int rowIdx = 1;
        for (Ticket t : tickets) {
            Row row = sheet.createRow(rowIdx++);
            row.createCell(0).setCellValue(t.getId());
            row.createCell(1).setCellValue(t.getTitulo());
        }
        wb.write(out);
        return out.toByteArray();
    }
}""")

    heading(doc, "Demo en Swagger UI", nivel=3)
    bullet(doc, "Login como ADMIN.")
    bullet(doc, "GET /api/dashboard/resumen → JSON con totales por estado.")
    bullet(doc, "GET /api/dashboard/tickets-por-tecnico → agrupado.")
    bullet(doc, "GET /api/ranking/categorias → categorías más reportadas con tiempo promedio.")
    bullet(doc, "GET /api/ranking/tecnicos/top/3 → top 3 técnicos.")
    bullet(doc, "GET /api/export/tickets/excel → descarga archivo .xlsx.")
    bullet(doc, "GET /api/export/tickets/pdf → descarga archivo .pdf.")
    bullet(doc, "Abrir el Excel/PDF descargado para que el docente vea el reporte completo.")

    heading(doc, "Posibles preguntas del docente", nivel=3)
    tabla_simple(doc,
                 ["Pregunta", "Respuesta sugerida"],
                 [
                     ["¿Por qué algunas queries son JPQL y otras nativas?",
                      "JPQL siempre que sea posible (portabilidad entre BDs). Nativas cuando usamos "
                      "funciones específicas de MySQL como TIMESTAMPDIFF dentro de AVG, porque Hibernate 6 "
                      "no infiere el tipo de retorno de FUNCTION() y rompe la validación."],
                     ["¿Qué es DTO y por qué no devuelven la entidad cruda?",
                      "DTO (Data Transfer Object) es un objeto plano para transferir datos. Para queries "
                      "agregadas no existe entidad correspondiente (COUNT, AVG no son atributos de Ticket), "
                      "así que mapeamos a un DTO específico. También evita exponer campos sensibles."],
                     ["¿Cómo se descarga un archivo desde un endpoint REST?",
                      "El controller devuelve byte[] con los headers Content-Type "
                      "(application/vnd.openxmlformats-...) y Content-Disposition (attachment; "
                      "filename=...). El navegador lo descarga automáticamente."],
                     ["¿Para qué el filtro de estado/fecha en el export?",
                      "Para reportes parciales: \"exportar solo tickets RESUELTOS de este mes\". El "
                      "controller recibe los query params, los pasa al service, y la query @Query filtra "
                      "con condiciones opcionales (IS NULL OR ...)."],
                 ],
                 anchos=[5, 11])

    saltopagina(doc)

    # ═══════════════════════════════════════════════════════════════
    # 7. ANEXO A — LEVANTAR PROYECTO
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "Anexo A — Cómo levantar el proyecto", nivel=1)

    parrafo(doc, "Prerrequisitos:", size=11, bold=True)
    bullet(doc, "JDK 21 (Eclipse Adoptium recomendado).")
    bullet(doc, "Maven 3.9+.")
    bullet(doc, "MySQL 8 corriendo en localhost:3306.")
    bullet(doc, "DBeaver o MySQL Workbench (para inspeccionar la BD).")
    bullet(doc, "VS Code con extensiones: Spring Boot Dashboard, Extension Pack for Java.")

    heading(doc, "Paso a paso", nivel=2)
    codigo(doc,
"""1. Clonar el repo
   git clone https://github.com/Carlos-Huane/Desarrollo-Web-Integrado_Back-End-.git
   cd Desarrollo-Web-Integrado_Back-End-

2. Crear la base de datos en MySQL
   CREATE DATABASE intranet_db CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci;

3. Configurar el password de MySQL en src/main/resources/application.properties
   spring.datasource.password=TU_PASSWORD

4. Levantar con Spring Boot Dashboard (boton play) o por terminal:
   mvn spring-boot:run

5. Verificar:
   - Logs: "Started IntranetApplication"
   - Swagger UI: http://localhost:8080/swagger-ui.html
   - DBeaver: intranet_db debe tener 7 tablas con datos""")

    callout(doc, "Si MySQL devuelve 'Public Key Retrieval is not allowed'",
            "Asegurarse de que el JDBC URL en application.properties tenga "
            "allowPublicKeyRetrieval=true. Es necesario porque MySQL 8 usa "
            "caching_sha2_password por defecto.")

    saltopagina(doc)

    # ═══════════════════════════════════════════════════════════════
    # 8. ANEXO B — CREDENCIALES
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "Anexo B — Credenciales de prueba", nivel=1)
    parrafo(doc,
            "DataInitializer.java siembra estos usuarios al primer arranque. Las contraseñas se guardan "
            "encriptadas con BCrypt, pero el login funciona con los valores en texto plano que aquí "
            "se muestran.")

    tabla_simple(doc,
                 ["Email", "Password", "Rol", "Nombre"],
                 [
                     ["admin@telecoperu.com", "admin123", "ADMIN", "Carlos Ramírez"],
                     ["gerente.ti@telecoperu.com", "admin123", "ADMIN", "Ana Torres"],
                     ["tecnico1@telecoperu.com", "tecnico123", "TECNICO", "Luis Mendoza"],
                     ["tecnico2@telecoperu.com", "tecnico123", "TECNICO", "Rosa Vega"],
                     ["pedro.sanchez@telecoperu.com", "cliente123", "CLIENTE", "Pedro Sánchez"],
                     ["maria.lopez@telecoperu.com", "cliente123", "CLIENTE", "María López"],
                     ["jose.quispe@telecoperu.com", "cliente123", "CLIENTE", "José Quispe"],
                 ],
                 anchos=[5, 3, 3, 5])

    saltopagina(doc)

    # ═══════════════════════════════════════════════════════════════
    # 9. ANEXO C — ENDPOINTS
    # ═══════════════════════════════════════════════════════════════
    heading(doc, "Anexo C — Listado completo de endpoints", nivel=1)

    heading(doc, "Auth (público)", nivel=2)
    tabla_simple(doc,
                 ["Método", "Ruta", "Descripción"],
                 [
                     ["POST", "/api/auth/login", "Inicia sesión y devuelve JWT"],
                     ["POST", "/api/auth/registro", "Registra un cliente (BindingResult)"],
                 ],
                 anchos=[2, 6, 8])

    heading(doc, "Usuarios", nivel=2)
    tabla_simple(doc,
                 ["Método", "Ruta", "Rol requerido"],
                 [
                     ["GET", "/api/usuarios", "ADMIN, TECNICO"],
                     ["GET", "/api/usuarios/activos", "ADMIN, TECNICO"],
                     ["GET", "/api/usuarios/tecnicos", "ADMIN, TECNICO"],
                     ["GET", "/api/usuarios/{id}", "ADMIN, TECNICO"],
                     ["POST", "/api/usuarios", "ADMIN"],
                     ["PUT", "/api/usuarios/{id}", "ADMIN"],
                     ["PATCH", "/api/usuarios/{id}/activar", "ADMIN"],
                     ["PATCH", "/api/usuarios/{id}/desactivar", "ADMIN"],
                 ],
                 anchos=[2, 7, 7])

    heading(doc, "Categorías", nivel=2)
    tabla_simple(doc,
                 ["Método", "Ruta", "Rol requerido"],
                 [
                     ["GET", "/api/categorias", "público"],
                     ["GET", "/api/categorias/todas", "público"],
                     ["POST", "/api/categorias", "ADMIN"],
                     ["PUT", "/api/categorias/{id}", "ADMIN"],
                     ["PATCH", "/api/categorias/{id}/activar", "ADMIN"],
                     ["PATCH", "/api/categorias/{id}/desactivar", "ADMIN"],
                     ["GET", "/api/categorias/{id}/subcategorias", "público"],
                     ["POST", "/api/categorias/subcategorias", "ADMIN"],
                 ],
                 anchos=[2, 7, 7])

    heading(doc, "Tickets", nivel=2)
    tabla_simple(doc,
                 ["Método", "Ruta", "Rol requerido"],
                 [
                     ["GET", "/api/tickets", "ADMIN, TECNICO"],
                     ["GET", "/api/tickets/{id}", "autenticado"],
                     ["GET", "/api/tickets/cliente/{clienteId}", "autenticado"],
                     ["GET", "/api/tickets/tecnico/{tecnicoId}", "autenticado"],
                     ["POST", "/api/tickets/cliente/{clienteId}", "autenticado"],
                     ["PATCH", "/api/tickets/{ticketId}/estado/{userId}", "ADMIN, TECNICO"],
                     ["GET", "/api/tickets/{ticketId}/historial", "autenticado"],
                 ],
                 anchos=[2, 7, 7])

    heading(doc, "Artículos (Base de Conocimiento)", nivel=2)
    tabla_simple(doc,
                 ["Método", "Ruta", "Rol requerido"],
                 [
                     ["GET", "/api/articulos", "autenticado"],
                     ["GET", "/api/articulos/todos", "autenticado"],
                     ["GET", "/api/articulos/{id}", "autenticado"],
                     ["GET", "/api/articulos/categoria/{categoriaId}", "autenticado"],
                     ["GET", "/api/articulos/buscar?q=texto", "autenticado"],
                     ["GET", "/api/articulos/top", "autenticado"],
                     ["POST", "/api/articulos/autor/{autorId}", "ADMIN, TECNICO"],
                     ["PUT", "/api/articulos/{id}", "ADMIN, TECNICO"],
                     ["PATCH", "/api/articulos/{id}/activar", "ADMIN, TECNICO"],
                     ["PATCH", "/api/articulos/{id}/desactivar", "ADMIN, TECNICO"],
                 ],
                 anchos=[2, 7, 7])

    heading(doc, "Dashboard / Ranking / Export / SLA", nivel=2)
    tabla_simple(doc,
                 ["Método", "Ruta", "Rol requerido"],
                 [
                     ["GET", "/api/dashboard/resumen", "ADMIN, TECNICO"],
                     ["GET", "/api/dashboard/tickets-por-estado", "ADMIN, TECNICO"],
                     ["GET", "/api/dashboard/tickets-por-prioridad", "ADMIN, TECNICO"],
                     ["GET", "/api/dashboard/tickets-por-categoria", "ADMIN, TECNICO"],
                     ["GET", "/api/dashboard/tickets-por-tecnico", "ADMIN, TECNICO"],
                     ["GET", "/api/ranking/categorias", "ADMIN, TECNICO"],
                     ["GET", "/api/ranking/categorias/top/{limite}", "ADMIN, TECNICO"],
                     ["GET", "/api/ranking/tecnicos", "ADMIN, TECNICO"],
                     ["GET", "/api/ranking/tecnicos/top/{limite}", "ADMIN, TECNICO"],
                     ["GET", "/api/export/tickets/excel", "ADMIN, TECNICO"],
                     ["GET", "/api/export/tickets/pdf", "ADMIN, TECNICO"],
                     ["GET", "/api/sla", "ADMIN"],
                     ["PUT", "/api/sla/{id}", "ADMIN"],
                 ],
                 anchos=[2, 8, 6])

    # ═══════════════════════════════════════════════════════════════
    # PIE FINAL
    # ═══════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()
    parrafo(doc, "— Fin de la guía —", size=10, italic=True, color=COLOR_GRIS,
            align=WD_ALIGN_PARAGRAPH.CENTER)
    parrafo(doc, "Grupo 6 · Desarrollo Web Integrado · 2026", size=9, italic=True, color=COLOR_GRIS,
            align=WD_ALIGN_PARAGRAPH.CENTER)

    output = "docs/Guia-Exposicion-Backend-Intranet.docx"
    doc.save(output)
    print(f"Generado: {output}")


if __name__ == "__main__":
    build_doc()
